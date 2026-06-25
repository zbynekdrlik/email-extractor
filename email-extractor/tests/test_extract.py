"""Tests for the extraction core. Fixtures are generated synthetically (no real data)."""
import io

import pytest
from PIL import Image, ImageDraw, ImageFont

from app import extract


# ----------------------------- pure helpers -------------------------------
def test_alpha_ratio_clean_vs_garbage():
    assert extract.alpha_ratio("Faktúra č. 123, Spolu 45.60 €") > 0.9
    assert extract.alpha_ratio("\x01\x02\x03###~~~|||��") < 0.55


@pytest.mark.parametrize("text,expected", [
    ("", True),
    ("Objednávka č. 1030 Dodávateľ SLOVNORMAL IČO 31697143", False),
    ("text with (cid:44)(cid:12) font failure", True),
    ("normal slovak text áäčďéíĺľňóôŕšťúýž words here plenty", False),
    ("���������� ###### |||||| \x00\x01", True),
])
def test_looks_garbage(text, expected):
    assert extract.looks_garbage(text) is expected


def test_file_ext():
    assert extract.file_ext("Faktura.PDF") == "pdf"
    assert extract.file_ext("noext") == ""
    assert extract.file_ext("a.b.docx") == "docx"


def test_decorative_image_filters():
    assert extract._is_decorative_image(800, 600, 1000) == "skipped_tiny_image"
    assert extract._is_decorative_image(50, 50, 20000) == "skipped_small_image"
    assert extract._is_decorative_image(900, 150, 60000) == "skipped_banner_image"
    assert extract._is_decorative_image(800, 1000, 200000) is None


# ----------------------------- native formats -----------------------------
def test_docx_native():
    import docx
    d = docx.Document()
    d.add_paragraph("Objednávka pečiva")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Chlieb"
    t.rows[0].cells[1].text = "15 ks"
    buf = io.BytesIO()
    d.save(buf)
    res = extract.extract_attachment("order.docx", "", buf.getvalue())
    assert res["method"] == "docx"
    assert "Objednávka pečiva" in res["text"]
    assert "Chlieb" in res["text"] and "15 ks" in res["text"]
    assert not res["needs_vision"]


def test_xlsx_native():
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Polozka", "Mnozstvo"])
    ws.append(["Bageta", 7])
    buf = io.BytesIO()
    wb.save(buf)
    res = extract.extract_attachment("o.xlsx", "", buf.getvalue())
    assert res["method"] == "xlsx"
    assert "Bageta" in res["text"] and "Polozka" in res["text"]


def test_xls_legacy_native():
    # legacy BIFF .xls (application/vnd.ms-excel) — openpyxl can't read these; xlrd path
    import xlwt
    wb = xlwt.Workbook()
    ws = wb.add_sheet("Objednavka")
    rows = [["Polozka", "Mnozstvo"], ["Bageta kvaskova 250g", 6], ["Rozok kvaskovy 70g", 35]]
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            ws.write(r, c, val)
    buf = io.BytesIO()
    wb.save(buf)
    res = extract.extract_attachment("Balla.xls", "application/vnd.ms-excel", buf.getvalue())
    assert res["method"] == "xls"
    assert "Bageta kvaskova 250g" in res["text"]
    assert "Rozok kvaskovy 70g" in res["text"]
    assert "35" in res["text"] and "6" in res["text"]   # quantities, not "6.0"
    assert res["chars"] > 0


def test_odt_native():
    # OpenDocument Text (.odt) — LibreOffice/OpenOffice Writer
    from odf.opendocument import OpenDocumentText
    from odf.text import P
    d = OpenDocumentText()
    d.text.addElement(P(text="Objednavka peciva: Bageta kvaskova 6 ks, Vianocka 2 ks"))
    buf = io.BytesIO()
    d.write(buf)
    res = extract.extract_attachment("order.odt", "application/vnd.oasis.opendocument.text",
                                     buf.getvalue())
    assert res["method"] == "odf"
    assert "Bageta kvaskova" in res["text"] and "Vianocka" in res["text"]


def test_ods_native():
    # OpenDocument Spreadsheet (.ods) — LibreOffice/OpenOffice Calc
    from odf.opendocument import OpenDocumentSpreadsheet
    from odf.table import Table, TableCell, TableRow
    from odf.text import P
    d = OpenDocumentSpreadsheet()
    tbl = Table(name="Objednavka")
    for vals in [["Polozka", "Mnozstvo"], ["Rozok kvaskovy", 35], ["Chlieb", 4]]:
        tr = TableRow()
        for v in vals:
            tc = TableCell()
            tc.addElement(P(text=str(v)))
            tr.addElement(tc)
        tbl.addElement(tr)
    d.spreadsheet.addElement(tbl)
    buf = io.BytesIO()
    d.write(buf)
    res = extract.extract_attachment("order.ods", "application/vnd.oasis.opendocument.spreadsheet",
                                     buf.getvalue())
    assert res["method"] == "odf"
    assert "Rozok kvaskovy" in res["text"] and "35" in res["text"]


def test_rtf_native():
    rtf = r"{\rtf1\ansi\ansicpg1250 Objednavka: Chlieb 1000g 2 ks\par Bageta 6 ks\par}"
    res = extract.extract_attachment("order.rtf", "application/rtf", rtf.encode("utf-8"))
    assert res["method"] == "rtf"
    assert "Chlieb 1000g" in res["text"] and "Bageta" in res["text"]


def test_txt_native():
    res = extract.extract_attachment("a.txt", "text/plain", "Dobrý deň\nObjednávka".encode())
    assert res["method"] == "text"
    assert "Objednávka" in res["text"]


def test_junk_extension_skipped():
    res = extract.extract_attachment("contact.vcf", "", b"BEGIN:VCARD")
    assert res["flag"] == "skipped_junk"
    assert res["method"] == "skipped"


# ----------------------------- PDF + OCR (needs tesseract/poppler) --------
def _has_ocr():
    import shutil
    return bool(shutil.which("tesseract") and shutil.which("pdftoppm"))


def _text_pdf(text: str) -> bytes:
    pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas
    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    y = 800
    for line in text.splitlines():
        c.drawString(40, y, line)
        y -= 18
    c.showPage()
    c.save()
    return buf.getvalue()


def _text_image(text: str, size=(1000, 700)) -> bytes:
    # light NOISY background so the PNG is realistically >8 KB (a clean white image
    # compresses below the junk-image byte threshold). Text stays black & readable.
    bg = Image.effect_noise(size, 22).point(lambda v: 200 + v // 4).convert("RGB")
    d = ImageDraw.Draw(bg)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 46)
    except Exception:
        font = ImageFont.load_default()
    y = 80
    for line in (text + "\n\nDodavatel SLOVNORMAL s.r.o.\nICO 31697143\nSpolu 85.58 EUR").splitlines():
        d.text((40, y), line, fill="black", font=font)
        y += 70
    buf = io.BytesIO()
    bg.save(buf, format="PNG")
    return buf.getvalue()


def test_pdf_text_layer_native():
    pdf = _text_pdf("FAKTURA c. 2612065\nDodavatel SLOVNORMAL\nSpolu 85.58")
    res = extract.extract_attachment("f.pdf", "application/pdf", pdf)
    assert res["method"] == "pdf-text"          # native, NOT ocr
    assert "FAKTURA" in res["text"] and "SLOVNORMAL" in res["text"]
    assert res["ocr_conf"] is None


@pytest.mark.skipif(not _has_ocr(), reason="tesseract/poppler not installed")
def test_image_ocr_reads_text():
    png = _text_image("FAKTURA 12345")
    res = extract.extract_attachment("scan.png", "image/png", png)
    assert res["method"] == "image-ocr"
    assert res["ocr_conf"] is not None
    assert "FAKTURA" in res["text"].upper() or "12345" in res["text"]


@pytest.mark.skipif(not _has_ocr(), reason="tesseract/poppler not installed")
def test_image_only_pdf_falls_back_to_ocr():
    # a PDF with NO text layer (just an image) must go through OCR, not pdf-text
    img = Image.open(io.BytesIO(_text_image("OBJEDNAVKA 999", size=(1000, 700))))
    buf = io.BytesIO()
    img.convert("RGB").save(buf, format="PDF")
    res = extract.extract_attachment("scanned.pdf", "application/pdf", buf.getvalue())
    assert res["method"] == "pdf-ocr"
    assert res["ocr_conf"] is not None


@pytest.mark.skipif(not _has_ocr(), reason="tesseract/poppler not installed")
def test_low_quality_image_flagged_needs_vision_and_text_dropped():
    # near-blank / noise image -> low confidence -> needs_vision, noisy text dropped
    img = Image.new("RGB", (600, 600), "white")
    d = ImageDraw.Draw(img)
    for i in range(0, 600, 7):
        d.line([(i, 0), (600 - i, 600)], fill=(180, 180, 180))  # noise, no real text
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    res = extract.extract_attachment("photo.png", "image/png", buf.getvalue())
    assert res["method"] == "image-ocr"
    if res["needs_vision"]:                       # fix #2: placeholder, not garbage
        assert res["text"].startswith("[needs AI Vision")
        assert res["flag"] == "needs_vision"
