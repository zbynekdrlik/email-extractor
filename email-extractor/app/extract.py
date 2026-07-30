"""
Attachment + body text extraction core.

Strategy (validated by the 100-email quality spike, 2026-06-25):
  - native text first: text-layer PDF (pdfplumber), docx, xlsx, txt/csv
  - PDF yields no text OR garbage (mojibake / "(cid:" / low alpha) -> OCR
  - images -> OCR (Tesseract ces+slk+eng, 300 DPI)
  - low OCR confidence on an image/scan -> needs_vision (route to AI Vision)

Two fixes over the spike:
  1. skip decorative signature/banner images (wide aspect ratio / tiny) instead
     of OCR-ing them into noise
  2. when needs_vision is set, DROP the noisy OCR text (keep a short placeholder)
     so ~KBs of garbage never pollute the combined text the classifier sees
"""
from __future__ import annotations

import io

OCR_LANG = "ces+slk+eng"
OCR_CONFIG = "--oem 1 --psm 6 -c preserve_interword_spaces=1"
OCR_DPI = 300
OCR_MAX_PAGES = 15
TEXT_PDF_MIN_CHARS_PER_PAGE = 50
GARBAGE_ALPHA = 0.55
NEEDS_VISION_CONF = 72
MIN_IMG_BYTES = 8000
MIN_IMG_PIXELS = 40000          # < ~200x200 = logo/icon
BANNER_ASPECT = 3.0            # wider/taller than this = decorative banner
JUNK_EXT = {"vcf", "ics", "p7s", "asc", "pgp", "gpg", "smime", "sig", "key"}
_ALPHA_EXTRA = set(".,;:-/€%()@")


def alpha_ratio(s: str) -> float:
    """Fraction of chars that are letters/digits/space/common punct (noise detector)."""
    if not s:
        return 0.0
    good = sum(1 for ch in s if ch.isalnum() or ch.isspace() or ch in _ALPHA_EXTRA)
    return round(good / len(s), 3)


def looks_garbage(text: str, alpha: float | None = None) -> bool:
    """True when a native text extraction is gibberish and should be OCR'd."""
    if not text:
        return True
    if "(cid:" in text:
        return True
    if text.count("�") / max(len(text), 1) > 0.02:
        return True
    if alpha is None:
        alpha = alpha_ratio(text)
    return alpha < GARBAGE_ALPHA


def _new_result(filename: str, mime: str, size: int) -> dict:
    return {"filename": filename or "(no name)", "mime": mime, "size": size,
            "method": None, "chars": 0, "ocr_conf": None, "pages": None,
            "alpha_ratio": None, "needs_vision": False, "native_garbage": False,
            "flag": "ok", "error": None, "text": ""}


def _set_text(res: dict, text: str):
    res["chars"] = len(text)
    res["alpha_ratio"] = alpha_ratio(text)
    res["text"] = text


def _ocr_images(images) -> tuple[str, float]:
    import pytesseract
    from pytesseract import Output
    texts, confs = [], []
    for img in images:
        if img.mode != "L":
            img = img.convert("L")
        data = pytesseract.image_to_data(img, lang=OCR_LANG, config=OCR_CONFIG,
                                          output_type=Output.DICT)
        words = []
        for w, conf in zip(data["text"], data["conf"], strict=False):
            try:
                ci = float(conf)
            except (TypeError, ValueError):
                ci = -1
            if w.strip() and ci >= 0:
                words.append(w)
                confs.append(ci)
        texts.append(" ".join(words))
    mean_conf = round(sum(confs) / len(confs), 1) if confs else 0.0
    return "\n".join(texts).strip(), mean_conf


def _is_decorative_image(width: int, height: int, size: int) -> str | None:
    """Return a skip-flag if the image is junk (logo/icon/banner), else None."""
    if size < MIN_IMG_BYTES:
        return "skipped_tiny_image"
    if (width * height) < MIN_IMG_PIXELS:
        return "skipped_small_image"
    if min(width, height) and max(width, height) / min(width, height) > BANNER_ASPECT:
        return "skipped_banner_image"   # fix #1: wide signature/marketing banners
    return None


def _apply_vision_gate(res: dict):
    """fix #2: when low-confidence OCR is flagged for vision, drop the noisy text."""
    if res["needs_vision"]:
        res["flag"] = "needs_vision"
        res["text"] = f"[needs AI Vision: {res['filename']}]"


def file_ext(filename: str) -> str:
    fn = (filename or "").strip().lower()
    return fn.rsplit(".", 1)[-1] if "." in fn else ""


def extract_attachment(filename: str, mime: str, data: bytes) -> dict:
    """Extract text from one attachment. Returns a result dict incl. `text`."""
    ext = file_ext(filename)
    mime = (mime or "").lower()
    res = _new_result((filename or "").strip(), mime, len(data))
    try:
        if ext in JUNK_EXT:
            res.update(method="skipped", flag="skipped_junk")
            return res

        if mime == "application/pdf" or ext == "pdf":
            return _extract_pdf(res, data)

        if mime.startswith("image/") or ext in {"png", "jpg", "jpeg", "tif",
                                                 "tiff", "bmp", "webp", "gif"}:
            return _extract_image(res, data)

        if ext == "docx" or "wordprocessingml" in mime:
            return _extract_docx(res, data)

        if ext in {"xlsx", "xlsm"} or "spreadsheetml" in mime:
            return _extract_xlsx(res, data)

        if ext == "xls" or (mime in {"application/vnd.ms-excel", "application/excel",
                                      "application/x-msexcel"} and ext not in {"csv", "txt"}):
            # Senders mislabel spreadsheets constantly, so the BYTES decide, not the
            # declared mime: an OOXML or CSV pushed into xlrd raised and the order text
            # was lost entirely (#23). Only real BIFF goes to xlrd.
            if data[:2] == b"PK":
                return _extract_xlsx(res, data)
            if not data.startswith(b"\xd0\xcf\x11\xe0"):
                res["method"] = "text"
                _set_text(res, data.decode("utf-8", errors="replace"))
                return res
            return _extract_xls(res, data)

        if ext in {"odt", "ods", "fodt", "fods"} or "opendocument" in mime:
            return _extract_odf(res, data)

        if ext == "doc" or mime == "application/msword":
            return _extract_doc(res, data)

        if ext == "rtf" or mime in {"application/rtf", "text/rtf"}:
            return _extract_rtf(res, data)

        if ext in {"txt", "csv", "log", "md"} or mime.startswith("text/"):
            res["method"] = "text"
            _set_text(res, data.decode("utf-8", errors="replace"))
            return res

        res.update(method="unsupported", flag="unsupported")
        return res
    except Exception as e:  # never let one attachment kill the email
        res.update(error=f"{type(e).__name__}: {e}", flag="error")
        return res


def _extract_pdf(res: dict, data: bytes) -> dict:
    import pdfplumber
    text, npages = "", 0
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            npages = len(pdf.pages)
            text = "\n".join((p.extract_text() or "") for p in pdf.pages).strip()
    except Exception as e:
        res["error"] = f"pdfplumber: {e}"
    res["pages"] = npages
    ypp = (len(text) / npages) if npages else 0
    ar = alpha_ratio(text)
    garbage = bool(text) and looks_garbage(text, ar)
    if text and ypp >= TEXT_PDF_MIN_CHARS_PER_PAGE and not garbage:
        res["method"] = "pdf-text"
        _set_text(res, text)
        return res
    # no text OR garbage -> OCR
    res["native_garbage"] = garbage
    from pdf2image import convert_from_bytes
    imgs = convert_from_bytes(data, dpi=OCR_DPI, first_page=1, last_page=OCR_MAX_PAGES)
    otext, conf = _ocr_images(imgs)
    res["method"] = "pdf-ocr"
    res["ocr_conf"] = conf
    _set_text(res, otext)
    res["needs_vision"] = (conf < NEEDS_VISION_CONF) or (len(otext) < 20)
    if not res["needs_vision"]:
        res["flag"] = ("garbage_native_ocr" if garbage
                       else ("ocr_truncated" if npages > OCR_MAX_PAGES else "ok"))
    _apply_vision_gate(res)
    return res


def _extract_image(res: dict, data: bytes) -> dict:
    from PIL import Image
    img = Image.open(io.BytesIO(data))
    skip = _is_decorative_image(img.width, img.height, len(data))
    if skip:
        res.update(method="skipped", flag=skip, pages=1)
        return res
    otext, conf = _ocr_images([img])
    res["method"] = "image-ocr"
    res["ocr_conf"] = conf
    res["pages"] = 1
    _set_text(res, otext)
    res["needs_vision"] = (conf < NEEDS_VISION_CONF) or (len(otext) < 10)
    _apply_vision_gate(res)
    return res


def _extract_docx(res: dict, data: bytes) -> dict:
    import docx
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.append("\t".join(c.text for c in row.cells))
    res["method"] = "docx"
    _set_text(res, "\n".join(parts))
    return res


def _fmt_cell(v) -> str:
    """One spreadsheet cell -> string. None -> "" (kept, for column alignment).
    Integer-valued floats -> int, so an order quantity 6.0 reads as "6", not "6.0"."""
    if v is None:
        return ""
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return str(v)


def _grid_rows(rows: list[list[str]]) -> list[str]:
    """Tab-join rows, padding each to the widest row so columns stay aligned.
    Empty cells are PRESERVED (a blank trailing 'quantity' column must line up so the
    AI can tell ordered rows from un-ordered ones); fully-empty rows are dropped."""
    ncols = max((len(r) for r in rows), default=0)
    out = []
    for r in rows:
        r = list(r) + [""] * (ncols - len(r))
        if any(c.strip() for c in r):
            out.append("\t".join(r))
    return out


def _extract_xlsx(res: dict, data: bytes) -> dict:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    chunks = []
    for ws in wb.worksheets:
        chunks.append(f"[Sheet: {ws.title}]")
        rows = [[_fmt_cell(c) for c in row] for row in ws.iter_rows(values_only=True)]
        chunks.extend(_grid_rows(rows))
    res["method"] = "xlsx"
    _set_text(res, "\n".join(chunks))
    return res


def _xls_cell(book, sheet, r: int, c: int) -> str:
    """One legacy .xls cell -> string, with date cells rendered as dates.

    xlrd hands date cells back as float serials, so a delivery date used to reach
    Postgres as e.g. `46237` and was simply lost (#23).
    """
    import xlrd
    cell = sheet.cell(r, c)
    if cell.ctype == xlrd.XL_CELL_DATE:
        try:
            dt = xlrd.xldate.xldate_as_datetime(cell.value, book.datemode)
        except (ValueError, OverflowError, xlrd.xldate.XLDateError):
            return _fmt_cell(cell.value)
        # Time-only cells (serial < 1) carry no meaningful date part.
        if dt.year == 1899 or (dt.hour or dt.minute) and cell.value < 1:
            return dt.strftime("%H:%M")
        return dt.strftime("%Y-%m-%d") if (dt.hour or dt.minute) == 0 \
            else dt.strftime("%Y-%m-%d %H:%M")
    return _fmt_cell(cell.value)


def _extract_xls(res: dict, data: bytes) -> dict:
    """Legacy BIFF .xls (application/vnd.ms-excel) — openpyxl can't read these."""
    import xlrd
    book = xlrd.open_workbook(file_contents=data)
    chunks = []
    for sh in book.sheets():
        chunks.append(f"[Sheet: {sh.name}]")
        rows = [[_xls_cell(book, sh, r, c) for c in range(sh.ncols)]
                for r in range(sh.nrows)]
        chunks.extend(_grid_rows(rows))
    res["method"] = "xls"
    _set_text(res, "\n".join(chunks))
    return res


ODF_MAX_REPEAT = 256   # Calc pads rows to 1024+ empty columns; keep the grid sane


def _odf_row_cells(row, table_mod, teletype) -> list[str]:
    """Cells of one ODF table row, with `number-columns-repeated` EXPANDED.

    Calc run-length-encodes runs of identical (usually empty) cells. Reading such a run
    as ONE cell shifted every following column LEFT, so the quantity was read out of the
    wrong column (#23).
    """
    out = []
    for c in row.getElementsByType(table_mod.TableCell):
        try:
            rep = int(c.getAttribute("numbercolumnsrepeated") or 1)
        except (TypeError, ValueError):
            rep = 1
        out.extend([teletype.extractText(c)] * max(1, min(rep, ODF_MAX_REPEAT)))
    while out and not out[-1].strip():
        out.pop()          # drop Calc's trailing padding, keep interior blanks
    return out


def _extract_flat_odf(res: dict, data: bytes) -> dict:
    """Flat ODF (.fods/.fodt): one XML document, no zip container."""
    from xml.etree import ElementTree as ET
    ns_t = "{urn:oasis:names:tc:opendocument:xmlns:table:1.0}"
    root = ET.fromstring(data.decode("utf-8", errors="replace"))
    parts = []
    for tbl in root.iter(f"{ns_t}table"):
        parts.append(f"[Sheet: {tbl.get(f'{ns_t}name') or ''}]")
        rows = []
        for row in tbl.iter(f"{ns_t}table-row"):
            cells = []
            for c in row.iter(f"{ns_t}table-cell"):
                try:
                    rep = int(c.get(f"{ns_t}number-columns-repeated") or 1)
                except (TypeError, ValueError):
                    rep = 1
                txt = "".join(c.itertext()).strip()
                cells.extend([txt] * max(1, min(rep, ODF_MAX_REPEAT)))
            while cells and not cells[-1].strip():
                cells.pop()
            rows.append(cells)
        parts.extend(_grid_rows(rows))
    if not parts:
        # Writer-flavoured flat ODF (.fodt): no tables, just paragraphs.
        parts = [t for t in ("".join(p.itertext()).strip() for p in root.iter(
            "{urn:oasis:names:tc:opendocument:xmlns:text:1.0}p")) if t]
    res["method"] = "odf-flat"
    _set_text(res, "\n".join(parts))
    return res


def _extract_odf(res: dict, data: bytes) -> dict:
    """OpenDocument .odt (Writer) / .ods (Calc) via odfpy."""
    from odf import table, teletype, text
    from odf.opendocument import load
    if not data[:2] == b"PK":
        # Flat single-XML ODF (.fods/.fodt): odfpy's load() only reads the zip-packaged
        # form and raised BadZipFile, which was swallowed into flag='error' with no text
        # at all (#23).
        return _extract_flat_odf(res, data)
    doc = load(io.BytesIO(data))
    mt = doc.mimetype or ""
    parts = []
    if "spreadsheet" in mt:
        for tbl in doc.getElementsByType(table.Table):
            parts.append(f"[Sheet: {tbl.getAttribute('name') or ''}]")
            rows = [_odf_row_cells(row, table, teletype)
                    for row in tbl.getElementsByType(table.TableRow)]
            parts.extend(_grid_rows(rows))
    else:
        for p in doc.getElementsByType(text.P):
            t = teletype.extractText(p)
            if t.strip():
                parts.append(t)
    res["method"] = "odf"
    _set_text(res, "\n".join(parts))
    return res


def _extract_doc(res: dict, data: bytes) -> dict:
    """Legacy binary .doc (application/msword) via the antiword system binary."""
    import subprocess
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".doc") as tf:
        tf.write(data)
        tf.flush()
        proc = subprocess.run(["antiword", tf.name], capture_output=True, timeout=60)
    txt = proc.stdout.decode("utf-8", errors="replace").strip()
    if not txt and proc.returncode != 0:
        err = proc.stderr.decode("utf-8", errors="replace")[:200]
        res.update(method="unsupported", flag="error", error=f"antiword rc={proc.returncode}: {err}")
        return res
    res["method"] = "doc"
    _set_text(res, txt)
    return res


def _extract_rtf(res: dict, data: bytes) -> dict:
    """Rich Text Format via striprtf (pure Python)."""
    from striprtf.striprtf import rtf_to_text
    txt = rtf_to_text(data.decode("utf-8", errors="replace"))
    res["method"] = "rtf"
    _set_text(res, txt)
    return res
